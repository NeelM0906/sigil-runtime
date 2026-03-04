"""Artifact generators for document and file creation.

Each generator takes structured content and produces a file on disk,
returning the path and metadata needed to register it as an artifact.
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path


def generate_pdf(title: str, content: str, author: str = "") -> bytes:
    """Generate a PDF document from text content. Returns PDF bytes."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    if author:
        pdf.set_font("Helvetica", "I", 11)
        pdf.cell(0, 8, f"Author: {author}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Body
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("## "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, stripped[3:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith("# "):
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 15)
            pdf.multi_cell(0, 10, stripped[2:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, f"  - {stripped[2:]}", new_x="LMARGIN", new_y="NEXT")
        elif stripped:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, stripped, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.ln(3)

    return pdf.output()


def generate_docx(title: str, content: str, author: str = "") -> bytes:
    """Generate a DOCX document from text content. Returns DOCX bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Document properties
    doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    # Title
    doc.add_heading(title, level=0)

    # Body
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
